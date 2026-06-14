import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image, ImageTk
import os
from datetime import datetime
import numpy as np

# ---- Import your custom modules ----
from model.landcover_model import process_image, CUSTOM_CLASSES
from utils.nlp_utils import generate_nlp_review

# ---- Define color palette for visualization ----
COLORS = {
    "Agriculture": (144, 238, 144),
    "Forest": (34, 139, 34),
    "Water": (0, 0, 255),
    "Urban": (255, 165, 0),
    "BarrenLand": (210, 180, 140)
}

def matrix_to_color_img(matrix, class_counts):
    """
    Convert a 2D class-name matrix to a color image for visualization.
    matrix: 2D numpy array with class names (strings).
    """
    color_img = np.zeros((matrix.shape[0], matrix.shape[1], 3), dtype=np.uint8)
    for cname, rgb in COLORS.items():
        color_img[matrix == cname] = rgb
    return Image.fromarray(color_img)

class SatelliteChangeDetector:
    def __init__(self, root):
        self.root = root
        self.root.title("Satellite Change Detector")
        self.image1_path = None
        self.image2_path = None
        self.analysis = None
        self.img1_imgtk = None
        self.img2_imgtk = None
        self.mat1_imgtk = None
        self.mat2_imgtk = None
        self.setup_ui()

    def setup_ui(self):
        main_frame = tk.Frame(self.root, padx=20, pady=20)
        main_frame.pack(fill="both", expand=True)

        # Images & Matrices
        img_matrix_frame = tk.Frame(main_frame)
        img_matrix_frame.pack()

        # Image 1 and its matrix
        left_frame_1 = tk.Frame(img_matrix_frame)
        left_frame_1.grid(row=0, column=0, padx=10, pady=5)
        tk.Label(left_frame_1, text="Image 1").pack()
        self.img1_label = tk.Label(left_frame_1, text="No Image", bg="gray")
        self.img1_label.pack(pady=3)
        self.matrix1_label = tk.Label(left_frame_1, text="Matrix 1", bg="white")
        self.matrix1_label.pack()

        # Image 2 and its matrix
        left_frame_2 = tk.Frame(img_matrix_frame)
        left_frame_2.grid(row=0, column=1, padx=10, pady=5)
        tk.Label(left_frame_2, text="Image 2").pack()
        self.img2_label = tk.Label(left_frame_2, text="No Image", bg="gray")
        self.img2_label.pack(pady=3)
        self.matrix2_label = tk.Label(left_frame_2, text="Matrix 2", bg="white")
        self.matrix2_label.pack()

        # Upload buttons
        upload_btn_frame = tk.Frame(main_frame)
        upload_btn_frame.pack(pady=10)
        tk.Button(upload_btn_frame, text="Upload Image 1", command=self.upload_img1).pack(side="left", padx=5)
        tk.Button(upload_btn_frame, text="Upload Image 2", command=self.upload_img2).pack(side="left", padx=5)

        # Action buttons
        action_btn_frame = tk.Frame(main_frame)
        action_btn_frame.pack(pady=15)
        tk.Button(action_btn_frame, text="Detect Change", command=self.detect_change, width=22, bg="#FF5722", fg="white").pack(side="left", padx=15)
        tk.Button(action_btn_frame, text="Generate Report", command=self.generate_report, width=22, bg="#2196F3", fg="white").pack(side="left", padx=15)

        # Results/Analysis
        self.text = tk.Text(main_frame, height=18, width=90, font=("Courier", 10), bg="white")
        self.text.pack(pady=10)

    def upload_img1(self):
        path = filedialog.askopenfilename(filetypes=[("Image files", "*.jpg *.jpeg *.png *.bmp")])
        if path:
            self.image1_path = path
            img = Image.open(path).resize((200, 150))
            self.img1_imgtk = ImageTk.PhotoImage(img)
            self.img1_label.config(image=self.img1_imgtk, text="")

            # Optionally clear the matrix display when a new image is loaded
            self.matrix1_label.config(image="", text="Matrix 1")

            self.text.insert('end', f"Image 1 loaded: {os.path.basename(path)}\n")

    def upload_img2(self):
        path = filedialog.askopenfilename(filetypes=[("Image files", "*.jpg *.jpeg *.png *.bmp")])
        if path:
            self.image2_path = path
            img = Image.open(path).resize((200, 150))
            self.img2_imgtk = ImageTk.PhotoImage(img)
            self.img2_label.config(image=self.img2_imgtk, text="")

            # Optionally clear the matrix display when a new image is loaded
            self.matrix2_label.config(image="", text="Matrix 2")

            self.text.insert('end', f"Image 2 loaded: {os.path.basename(path)}\n")

    def detect_change(self):
        if not self.image1_path or not self.image2_path:
            messagebox.showerror("Error", "Upload both images")
            return

        self.text.insert('end', "Analyzing...\n")
        self.text.see('end')

        res1 = process_image(self.image1_path)
        res2 = process_image(self.image2_path)
        nlp = generate_nlp_review(res1['dominant_class'], res2['dominant_class'], res1['class_counts'], res2['class_counts'])
        self.analysis = dict(res1=res1, res2=res2, nlp=nlp)

        # Show matrix as color mask (resize for display)
        mat1_img = matrix_to_color_img(res1['matrix'], res1['class_counts']).resize((200, 150), Image.NEAREST)
        mat2_img = matrix_to_color_img(res2['matrix'], res2['class_counts']).resize((200, 150), Image.NEAREST)
        self.mat1_imgtk = ImageTk.PhotoImage(mat1_img)
        self.mat2_imgtk = ImageTk.PhotoImage(mat2_img)
        self.matrix1_label.config(image=self.mat1_imgtk, text="")
        self.matrix2_label.config(image=self.mat2_imgtk, text="")

        self.text.insert('end', f"Image 1 Dominant: {res1['dominant_class']} | Image 2 Dominant: {res2['dominant_class']}\n")
        self.text.insert('end', f"Change Analysis:\n{nlp}\n\n")
        self.text.see('end')

    def generate_report(self):
        if not self.analysis:
            messagebox.showerror("Error", "Run analysis first!")
            return
        from docx import Document
        doc = Document()
        doc.add_heading('Satellite Landscape Change Detection Report', 0)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Image 1: {os.path.basename(self.image1_path)}")
        doc.add_paragraph(f"Image 2: {os.path.basename(self.image2_path)}")
        doc.add_heading('Analysis Summary', level=1)
        doc.add_paragraph(f"Image 1 Dominant: {self.analysis['res1']['dominant_class']}")
        doc.add_paragraph(f"Image 2 Dominant: {self.analysis['res2']['dominant_class']}")
        doc.add_paragraph(f"Class Distribution Image 1: {self.analysis['res1']['class_counts']}")
        doc.add_paragraph(f"Class Distribution Image 2: {self.analysis['res2']['class_counts']}")
        doc.add_heading('NLP Review & Recommendations', level=1)
        doc.add_paragraph(self.analysis['nlp'])
        save_path = filedialog.asksaveasfilename(defaultextension=".docx",
                            filetypes=[("Word Document", "*.docx")],
                            initialfile="Satellite_Change_Report.docx")
        if save_path:
            doc.save(save_path)
            messagebox.showinfo("Saved", f"Report saved to {save_path}")

if __name__ == "__main__":
    root = tk.Tk()
    app = SatelliteChangeDetector(root)
    root.mainloop()
